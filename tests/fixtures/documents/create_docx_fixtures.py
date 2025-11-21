# tests/fixtures/documents/create_docx_fixtures.py
# Script to create DOCX test fixtures w/ formatting for testing

from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# * Create basic DOCX resume w/ formatting for testing
def create_basic_formatted_resume():
    doc = Document()

    # header w/ bold name
    header = doc.add_paragraph()
    name_run = header.add_run("John Doe")
    name_run.bold = True
    name_run.font.size = 180000  # 18pt
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # title
    title = doc.add_paragraph("Software Engineer")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # contact info
    contact = doc.add_paragraph(
        "john.doe@email.com | (555) 123-4567 | linkedin.com/in/johndoe"
    )
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # empty line
    doc.add_paragraph()

    # professional summary section
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
    summary_run.bold = True

    doc.add_paragraph(
        "Experienced software engineer w/ 5+ years developing web applications using modern frameworks & technologies."
    )

    # empty line
    doc.add_paragraph()

    # skills section w/ formatting
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run("TECHNICAL SKILLS")
    skills_run.bold = True

    # bullet points w/ mixed formatting
    skill1 = doc.add_paragraph("• Programming Languages: ")
    skill1_bold = skill1.add_run("Python, JavaScript, TypeScript")
    skill1_bold.bold = True

    doc.add_paragraph("• Frameworks: React, Django, FastAPI")

    skill3 = doc.add_paragraph("• Tools: ")
    skill3_italic = skill3.add_run("Git, Docker, AWS")
    skill3_italic.italic = True

    # empty line
    doc.add_paragraph()

    # experience section
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run("PROFESSIONAL EXPERIENCE")
    exp_run.bold = True

    # job 1
    job1_title = doc.add_paragraph()
    job1_title.add_run("Senior Software Engineer").bold = True
    job1_title.add_run(" | Tech Corp | Jan 2021 - Present")

    doc.add_paragraph(
        "• Developed & maintained microservices architecture serving 500K+ users"
    )
    doc.add_paragraph("• Led code reviews & mentored junior developers")

    # job 2
    job2_title = doc.add_paragraph()
    job2_title.add_run("Software Engineer").bold = True
    job2_title.add_run(" | StartupCo | Jun 2019 - Dec 2020")

    doc.add_paragraph("• Built full-stack web applications using React & Python")
    doc.add_paragraph("• Collaborated w/ product team to define technical requirements")

    return doc


# * Create simple DOCX w/ basic formatting for testing
def create_simple_formatted_doc():
    doc = Document()

    # title
    title = doc.add_paragraph()
    title_run = title.add_run("Simple Document")
    title_run.bold = True
    title_run.font.size = 160000  # 16pt

    # paragraph w/ mixed formatting
    para = doc.add_paragraph("This is a paragraph with ")
    bold_run = para.add_run("bold text")
    bold_run.bold = True
    para.add_run(" and ")
    italic_run = para.add_run("italic text")
    italic_run.italic = True
    para.add_run(" and normal text.")

    # bullet list
    doc.add_paragraph("• First item")
    doc.add_paragraph("• Second item")
    doc.add_paragraph("• Third item")

    return doc


# * Create DOCX w/ edge cases for testing
def create_edge_case_doc():
    doc = Document()

    # document w/ only whitespace lines
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Single line")
    doc.add_paragraph("")
    doc.add_paragraph("Another line")
    doc.add_paragraph("")

    return doc


if __name__ == "__main__":
    fixture_dir = Path(__file__).parent

    # create basic formatted resume
    basic_doc = create_basic_formatted_resume()
    basic_doc.save(str(fixture_dir / "basic_formatted_resume.docx"))
    print("Created basic_formatted_resume.docx")

    # create simple formatted document
    simple_doc = create_simple_formatted_doc()
    simple_doc.save(str(fixture_dir / "simple_formatted.docx"))
    print("Created simple_formatted.docx")

    # create edge case document
    edge_doc = create_edge_case_doc()
    edge_doc.save(str(fixture_dir / "edge_cases.docx"))
    print("Created edge_cases.docx")
