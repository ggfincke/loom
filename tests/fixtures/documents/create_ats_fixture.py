#!/usr/bin/env python3
# tests/fixtures/documents/create_ats_fixture.py
# Script to create a DOCX file w/ common ATS-problematic structures for testing

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Create a resume w/ multiple ATS compatibility issues.
def create_ats_problem_resume():
    doc = Document()

    # 1. Add content to header (ATS often skips headers)
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "John Doe | Software Engineer | john@email.com"

    # 2. Add content to footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page 1 | Updated December 2024"

    # 3. Add a table (very common ATS issue)
    doc.add_heading("Contact Information", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Email:"
    table.cell(0, 1).text = "john.doe@email.com"
    table.cell(1, 0).text = "Phone:"
    table.cell(1, 1).text = "(555) 123-4567"

    # 4. Add regular content
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Experienced software engineer with 5+ years of experience in Python, "
        "JavaScript, and cloud technologies."
    )

    # 5. Add another table for skills (common pattern that breaks ATS)
    doc.add_heading("Skills", level=1)
    skills_table = doc.add_table(rows=2, cols=3)
    skills_table.cell(0, 0).text = "Python"
    skills_table.cell(0, 1).text = "JavaScript"
    skills_table.cell(0, 2).text = "TypeScript"
    skills_table.cell(1, 0).text = "AWS"
    skills_table.cell(1, 1).text = "Docker"
    skills_table.cell(1, 2).text = "PostgreSQL"

    # 6. Add experience section
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Senior Software Engineer at TechCorp (2020-Present)")
    doc.add_paragraph("- Developed microservices architecture")
    doc.add_paragraph("- Led team of 5 engineers")

    # 7. Add education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("BS Computer Science, State University (2015)")

    # save
    output_path = "ats_problem_resume.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    print("ATS issues included:")
    print("  - Header with contact info")
    print("  - Footer with page info")
    print("  - 2 tables (contact info and skills)")

    return output_path


if __name__ == "__main__":
    create_ats_problem_resume()
