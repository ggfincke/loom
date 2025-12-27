# src/core/ats_analyzer.py
# Core ATS (Applicant Tracking System) compatibility analyzer
# Performs structural analysis of DOCX files to detect ATS-problematic elements

from __future__ import annotations

from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document


# issue severity levels for ATS compatibility
class Severity(Enum):
    CRITICAL = "critical"  # ATS parsing blockers (tables, columns, text boxes)
    WARNING = "warning"  # Likely parsing degradation (headers, images)
    INFO = "info"  # Best practices (section naming, formatting)


# categories of ATS compatibility issues
class IssueCategory(Enum):
    TABLE = "table"
    COLUMNS = "columns"
    TEXTBOX = "textbox"
    HEADER_FOOTER = "header_footer"
    IMAGE = "image"
    DRAWING = "drawing"
    CONTACT = "contact"
    SECTION_HEADERS = "section_headers"
    BULLETS = "bullets"
    UNICODE = "unicode"
    DATES = "dates"


# single ATS compatibility issue
@dataclass
class ATSIssue:
    category: IssueCategory
    severity: Severity
    description: str
    location: Optional[str] = None  # e.g., "Line 15" or "Section 2"
    suggestion: Optional[str] = None  # How to fix

    # convert to JSON-serializable dict
    def to_dict(self) -> Dict[str, Any]:
        return {
            "category": self.category.value,
            "severity": self.severity.value,
            "description": self.description,
            "location": self.location,
            "suggestion": self.suggestion,
        }


# complete ATS compatibility report
@dataclass
class ATSReport:
    score: int  # 0-100, higher = better
    issues: List[ATSIssue] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    file_type: str = "docx"  # 'docx' or 'tex'

    @property
    def critical_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == Severity.CRITICAL)

    @property
    def warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == Severity.WARNING)

    @property
    def info_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == Severity.INFO)

    # resume passes if no critical issues
    @property
    def passed(self) -> bool:
        return self.critical_count == 0

    # get issues filtered by severity
    def issues_by_severity(self, severity: Severity) -> List[ATSIssue]:
        return [i for i in self.issues if i.severity == severity]

    # convert to JSON-serializable dict
    def to_dict(self) -> Dict[str, Any]:
        return {
            "version": 1,
            "score": self.score,
            "passed": self.passed,
            "file_type": self.file_type,
            "summary": {
                "critical": self.critical_count,
                "warning": self.warning_count,
                "info": self.info_count,
                "total": len(self.issues),
            },
            "issues": [issue.to_dict() for issue in self.issues],
            "recommendations": self.recommendations,
        }


# * Calculate ATS compatibility score (0-100, higher = better) w/ severity-based penalties
def calculate_score(issues: List[ATSIssue]) -> int:
    score = 100

    penalties = {
        Severity.CRITICAL: 25,  # Tables, columns = severe penalties
        Severity.WARNING: 10,
        Severity.INFO: 2,
    }

    for issue in issues:
        score -= penalties[issue.severity]

    return max(0, score)


# * Generate actionable recommendations based on detected issues
def generate_recommendations(issues: List[ATSIssue]) -> List[str]:
    recommendations: List[str] = []
    categories_seen = {issue.category for issue in issues}

    if IssueCategory.TABLE in categories_seen:
        recommendations.append(
            "Remove all tables & use plain paragraphs with consistent spacing"
        )

    if IssueCategory.COLUMNS in categories_seen:
        recommendations.append(
            "Convert to single-column layout to ensure correct text ordering"
        )

    if IssueCategory.TEXTBOX in categories_seen:
        recommendations.append(
            "Remove text boxes & place content in regular paragraphs"
        )

    if IssueCategory.HEADER_FOOTER in categories_seen:
        recommendations.append(
            "Move header/footer content to first lines of document body"
        )

    if (
        IssueCategory.IMAGE in categories_seen
        or IssueCategory.DRAWING in categories_seen
    ):
        recommendations.append(
            "Remove decorative images; ensure no critical text is embedded in graphics"
        )

    if IssueCategory.BULLETS in categories_seen:
        recommendations.append("Use standard bullet characters throughout (-, *, or •)")

    if IssueCategory.UNICODE in categories_seen:
        recommendations.append(
            "Replace unusual unicode characters with standard ASCII equivalents"
        )

    if IssueCategory.SECTION_HEADERS in categories_seen:
        recommendations.append(
            "Use standard section names: Experience, Education, Skills, Summary"
        )

    if IssueCategory.CONTACT in categories_seen:
        recommendations.append(
            "Ensure email & phone are on separate lines in plain text format"
        )

    if IssueCategory.DATES in categories_seen:
        recommendations.append(
            "Use consistent date format throughout (e.g., 'January 2020' or '01/2020')"
        )

    return recommendations


# Word XML namespace
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# * Analyze DOCX file for structural ATS issues (no AI required)
# Detects: tables, multi-column layouts, text boxes, headers/footers w/ content, inline images/shapes, drawing elements
def analyze_docx_structure(doc_path: Path) -> List[ATSIssue]:
    doc = Document(str(doc_path))
    issues: List[ATSIssue] = []

    # 1. Tables - CRITICAL
    if doc.tables:
        issues.append(
            ATSIssue(
                category=IssueCategory.TABLE,
                severity=Severity.CRITICAL,
                description=f"{len(doc.tables)} table(s) detected - ATS may fail to parse content",
                suggestion="Convert table content to plain paragraphs",
            )
        )

    # 2. Headers/Footers w/ content - WARNING
    for i, section in enumerate(doc.sections):
        # check header
        header_text = "".join(p.text for p in section.header.paragraphs).strip()
        if header_text:
            # truncate long header text for display
            preview = header_text[:50] + "..." if len(header_text) > 50 else header_text
            issues.append(
                ATSIssue(
                    category=IssueCategory.HEADER_FOOTER,
                    severity=Severity.WARNING,
                    description=f"Header contains text that ATS may skip: '{preview}'",
                    location=f"Section {i + 1}" if len(doc.sections) > 1 else None,
                    suggestion="Move contact info to document body",
                )
            )

        # check footer
        footer_text = "".join(p.text for p in section.footer.paragraphs).strip()
        if footer_text:
            preview = footer_text[:50] + "..." if len(footer_text) > 50 else footer_text
            issues.append(
                ATSIssue(
                    category=IssueCategory.HEADER_FOOTER,
                    severity=Severity.WARNING,
                    description=f"Footer contains text that ATS may skip: '{preview}'",
                    location=f"Section {i + 1}" if len(doc.sections) > 1 else None,
                    suggestion="Move footer content to document body",
                )
            )

    # 3. Multi-column layouts - CRITICAL
    for i, section in enumerate(doc.sections):
        sect_pr = section._sectPr
        if sect_pr is not None:
            cols = sect_pr.find(f".//{W_NS}cols")
            if cols is not None:
                num_cols = cols.get(f"{W_NS}num", "1")
                try:
                    if int(num_cols) > 1:
                        issues.append(
                            ATSIssue(
                                category=IssueCategory.COLUMNS,
                                severity=Severity.CRITICAL,
                                description=f"{num_cols}-column layout detected - text ordering will be scrambled",
                                location=(
                                    f"Section {i + 1}"
                                    if len(doc.sections) > 1
                                    else None
                                ),
                                suggestion="Use single-column format",
                            )
                        )
                except ValueError:
                    pass  # ignore non-numeric values

    # 4. Text boxes - CRITICAL
    body = doc.element.body
    textboxes = body.findall(f".//{W_NS}txbxContent")
    if textboxes:
        issues.append(
            ATSIssue(
                category=IssueCategory.TEXTBOX,
                severity=Severity.CRITICAL,
                description=f"{len(textboxes)} text box(es) detected - content often invisible to ATS",
                suggestion="Replace text boxes with regular paragraphs",
            )
        )

    # 5. Drawings/shapes - WARNING
    drawings = body.findall(f".//{W_NS}drawing")
    picts = body.findall(f".//{W_NS}pict")
    total_drawings = len(drawings) + len(picts)
    if total_drawings:
        issues.append(
            ATSIssue(
                category=IssueCategory.DRAWING,
                severity=Severity.WARNING,
                description=f"{total_drawings} drawing/shape element(s) detected",
                suggestion="Remove decorative graphics; ensure no text is in images",
            )
        )

    # 6. Inline shapes (images, SmartArt, etc.) - WARNING
    if doc.inline_shapes:
        issues.append(
            ATSIssue(
                category=IssueCategory.IMAGE,
                severity=Severity.WARNING,
                description=f"{len(doc.inline_shapes)} inline image(s)/shape(s) detected",
                suggestion="Remove images or ensure they are purely decorative",
            )
        )

    return issues


# * Analyze resume for ATS compatibility & return report w/ score, issues, & recommendations
def analyze_resume_ats(
    doc_path: Path,
    ai_issues: Optional[List[ATSIssue]] = None,
) -> ATSReport:
    # structural analysis
    issues = analyze_docx_structure(doc_path)

    # merge AI issues if provided
    if ai_issues:
        issues.extend(ai_issues)

    # calculate score
    score = calculate_score(issues)

    # generate recommendations
    recommendations = generate_recommendations(issues)

    return ATSReport(
        score=score,
        issues=issues,
        recommendations=recommendations,
        file_type="docx",
    )
