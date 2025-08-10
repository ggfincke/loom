from pathlib import Path
from docx import Document
import json

# read docx file & return text content
def read_docx(path: Path) -> dict[int, str]:
    doc = Document(str(path))
    lines = {}
    line_number = 1

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines[line_number] = text
            line_number += 1

    return lines

# write text to docx file
def write_docx(lines: dict[int, str], output_path: Path) -> None:
    doc = Document()
    for line_num in sorted(lines.keys()):
        doc.add_paragraph(lines[line_num])
    doc.save(str(output_path))

# number lines in a resume
def number_lines(resume: dict[int, str]) -> str:
    return "\n".join(f"{i:>4} {text}" for i, text in sorted(resume.items()))

# read text from file
def read_text(path: Path) -> str:
    return Path(path).read_text(encoding="utf-8")

# write text to JSON file 
def write_json(obj: dict, path: Path) -> None:
    path.write_text(json.dumps(obj, indent=2), encoding="utf-8")