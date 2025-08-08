import os
from pathlib import Path
from typing import Optional

import typer
import json

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI

app = typer.Typer(help="Tailor resumes using the OpenAI Responses API")

# read docx file & return text content
def read_docx(path: Path) -> dict[int, str]:
    doc = Document(str(path))
    lines = {}
    line_number = 1

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines[line_number] = text
            line_number += 1

    return lines

# write text to docx file
def write_docx(lines: dict[int, str], output_path: Path) -> None:
    doc = Document()
    for line_num in sorted(lines.keys()):
        doc.add_paragraph(lines[line_num])
    doc.save(str(output_path))

# build sectionizer prompt for LLM
def build_sectionizer_prompt(resume_with_line_numbers: str) -> str:
    return (
        "You are a resume section parser. Your task is to analyze a resume provided as "
        "numbered lines and return a strict JSON object describing sections and their line ranges.\n\n"
        "Rules:\n"
        "1) Use 1-based line numbers that exactly match the provided numbering.\n"
        "2) Detect common headings and variants (e.g., 'PROFESSIONAL SUMMARY', 'ABOUT ME' → SUMMARY; "
        "'TECHNOLOGIES', 'TOOLS' → SKILLS; 'WORK EXPERIENCE', 'EMPLOYMENT' → EXPERIENCE; "
        "'PERSONAL PROJECTS' → PROJECTS; 'EDUCATION'/'ACADEMICS' → EDUCATION).\n"
        "3) For each section, return start_line and end_line inclusive. Headings belong to their section.\n"
        "4) Include a 'confidence' score in [0,1] for each section.\n"
        "5) If you can identify repeated substructures (like individual experience entries), include them under 'subsections' with start/end lines and basic metadata if visible (company/title/date_range/location). Omit fields you cannot infer.\n"
        "6) Output ONLY JSON. No prose.\n\n"
        "JSON schema to follow exactly:\n"
        "{\n"
        "  \"sections\": [\n"
        "    {\n"
        "      \"name\": \"SUMMARY|SKILLS|EXPERIENCE|PROJECTS|EDUCATION|OTHER\",\n"
        "      \"heading_text\": \"<exact heading line text>\",\n"
        "      \"start_line\": <int>,\n"
        "      \"end_line\": <int>,\n"
        "      \"confidence\": <float>,\n"
        "      \"subsections\": [\n"
        "        {\n"
        "          \"name\": \"EXPERIENCE_ITEM|PROJECT_ITEM|EDUCATION_ITEM\",\n"
        "          \"start_line\": <int>,\n"
        "          \"end_line\": <int>,\n"
        "          \"meta\": {\"company\"?: \"string\", \"title\"?: \"string\", \"date_range\"?: \"string\", \"location\"?: \"string\"}\n"
        "        }\n"
        "      ]\n"
        "    }\n"
        "  ],\n"
        "  \"normalized_order\": [\"...\"],\n"
        "  \"notes\": \"string\"\n"
        "}\n\n"
        "Resume (numbered lines start at 1):\n"
        f"{resume_with_line_numbers}\n"
    )

# build tailoring prompt for LLM
def build_tailor_prompt(job_info: str,
                        resume_with_line_numbers: str,
                        sections_json: str | None = None) -> str:
    return (
        "You are a resume editor tasked with tailoring a resume (provided as numbered lines) "
        "to a specific job description. Return a STRICT JSON object with surgical edits by line number.\n\n"
        "Objectives:\n"
        "- Prioritize relevance to the job description.\n"
        "- Keep truthful content only; do not invent experience.\n"
        "- Prefer quantified impact and concise bullets.\n"
        "- Keep the existing visual style (single column, headings, bullets). Only change text.\n\n"
        "Line-numbering rules:\n"
        "1) Use the exact 1-based line numbers provided.\n"
        "2) For small tweaks, use 'replace_line'. For multi-line rewrites, use 'replace_range'.\n"
        "3) To add a new bullet directly after a line, use 'insert_after'.\n"
        "4) To remove irrelevant content, use 'delete_range'.\n"
        "5) Never output lines that don't exist; validate with current_snippet to avoid drift.\n\n"
        "Output ONLY JSON matching this schema:\n"
        "{\n"
        "  \"edits\": [\n"
        "    {\"op\": \"replace_line\", \"line\": <int>, \"current_snippet\": \"string\", \"replacement\": \"string\"},\n"
        "    {\"op\": \"replace_range\", \"start_line\": <int>, \"end_line\": <int>, \"current_snippet\": \"string\", \"replacement_lines\": [\"string\", \"...\"]},\n"
        "    {\"op\": \"insert_after\", \"line\": <int>, \"new_lines\": [\"string\", \"...\"]},\n"
        "    {\"op\": \"delete_range\", \"start_line\": <int>, \"end_line\": <int>, \"reason\": \"string\"}\n"
        "  ],\n"
        "  \"rationale\": \"string\"\n"
        "}\n\n"
        "Job Description:\n"
        f"{job_info}\n\n"
        + (
            f"Known Sections (JSON):\n{sections_json}\n\n" if sections_json else ""
        ) +
        "Resume (numbered lines start at 1):\n"
        f"{resume_with_line_numbers}\n"
    )

# helper functions

# number lines in a resume
def number_lines(resume: dict[int, str]) -> str:
    return "\n".join(f"{i:>4} {text}" for i, text in sorted(resume.items()))

# validate JSON response from OpenAI API
def openai_json(prompt: str, model: str = "gpt-4o-mini") -> dict:
    load_dotenv()
    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("Missing OPENAI_API_KEY in environment or .env")
    client = OpenAI()
    resp = client.responses.create(model=model, input=prompt, temperature=0.2)
    # ensure valid JSON (model should already be constrained by prompt)
    try:
        return json.loads(resp.output_text)
    except json.JSONDecodeError as e:
        # fail
        raise RuntimeError(f"Model did not return valid JSON. Raw:\n{resp.output_text}") from e

def read_text(path: Path) -> str:
    return Path(path).read_text(encoding="utf-8")

# CLI commands
@app.command()
def sectionize(
    resume_path: Path = typer.Argument(..., exists=True, readable=True, help="Path to source resume .docx"),
    out_json: Path = typer.Option(Path("sections.json"), help="Where to write the sections JSON"),
    model: str = typer.Option("gpt-4o-mini", help="OpenAI model name"),
):
    '''
    Parse a resume (.docx) into sections using OpenAI's Response API
    '''
    lines = read_docx(resume_path)
    numbered = number_lines(lines)
    prompt = build_sectionizer_prompt(numbered)
    data = openai_json(prompt, model=model)
    out_json.write_text(json.dumps(data, indent=2), encoding="utf-8")
    typer.echo(f"Wrote {out_json}")

@app.command()
def tailer(
    job_info: Path = typer.Argument(..., help="Job description text to tailor resume for"),
    resume_path: Path = typer.Argument(..., exists=True, readable=True, help="Path to source resume .docx"),
    sections_path: Optional[Path] = typer.Option(None, help="Optional sections.json from the 'sections' command"),
    out_json: Path = typer.Option(Path("edits.json"), help="Where to write the edits JSON"),
    model: str = typer.Option("gpt-4o-mini", help="OpenAI model name"),
):
    '''
    Tailor a resume (.docx) to a job description using OpenAI's Response API
    Generates a JSON object with edits by line number
    '''
    job_text = read_text(job_info)
    lines = read_docx(resume_path)
    numbered = number_lines(lines)
    sections_json_str = None
    if sections_path and sections_path.exists():
        sections_json_str = sections_path.read_text(encoding="utf-8")

    prompt = build_tailor_prompt(job_text, numbered, sections_json_str)
    data = openai_json(prompt, model=model)
    out_json.write_text(json.dumps(data, indent=2), encoding="utf-8")
    typer.echo(f"Wrote {out_json}")

if __name__ == "__main__":
    app()
