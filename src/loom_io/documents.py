# src/loom_io/documents.py
# Document I/O operations for reading & writing DOCX files w/ formatting preservation, plus basic LaTeX/text support

import zipfile
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml import OxmlElement
from typing import Dict, Tuple, Any, List, Set

from ..core.types import Lines
from ..core.exceptions import LaTeXError, TypstError, DocumentParseError
from ..core.verbose import vlog_file_read, vlog_file_write
from .generics import ensure_parent
from .latex_patterns import is_preservable_content, requires_trailing_blank

# ! Related: LaTeX patterns defined in latex_patterns.py
# ! Moved to lazy import below to avoid circular dependency w/ core/validation


# * Read DOCX file & return text content w/ document object
def read_docx_with_formatting(path: Path) -> Tuple[Lines, Any, Dict[int, Paragraph]]:
    try:
        doc = Document(str(path))
    except FileNotFoundError:
        raise DocumentParseError(f"DOCX file not found: {path}")
    except zipfile.BadZipFile:
        raise DocumentParseError(f"Invalid DOCX file (not a valid zip archive): {path}")
    except Exception as e:
        raise DocumentParseError(f"Failed to open DOCX file {path}: {e}") from e

    lines = {}
    paragraph_map = {}
    line_number = 1

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines[line_number] = text
            paragraph_map[line_number] = p
            line_number += 1

    vlog_file_read(path, path.stat().st_size if path.exists() else None)
    return lines, doc, paragraph_map


# * Read DOCX file & return text content (backward compatibility)
def read_docx(path: Path) -> Lines:
    lines, _, _ = read_docx_with_formatting(path)
    return lines


# * Read LaTeX (.tex) file as numbered lines
def read_latex(path: Path, preserve_structure: bool = False) -> Lines:
    # ! Lazy import to avoid circular dependency
    from .latex_handler import validate_basic_latex_syntax

    try:
        text = Path(path).read_text(encoding="utf-8")
        vlog_file_read(path, len(text))
    except UnicodeDecodeError as e:
        raise LaTeXError(f"Cannot decode LaTeX file {path}: {e}")
    except Exception as e:
        raise LaTeXError(f"Cannot read LaTeX file {path}: {e}")

    # Validate basic LaTeX syntax
    if not validate_basic_latex_syntax(text):
        raise LaTeXError(f"Invalid LaTeX syntax detected in {path}")

    lines: Lines = {}
    line_number = 1

    if preserve_structure:
        # Structured mode: preserve comments, commands, strategic empty lines
        for raw in text.splitlines():
            t = raw.strip()

            # Preserve important LaTeX constructs
            if is_preservable_content(t):
                lines[line_number] = t
                line_number += 1
            # Preserve strategic empty lines after structural commands
            elif raw == "" and line_number > 1:
                prev_line = lines.get(line_number - 1, "")
                if requires_trailing_blank(prev_line):
                    lines[line_number] = ""
                    line_number += 1
    else:
        # Basic mode: strip whitespace, keep only non-empty content
        for raw in text.splitlines():
            t = raw.strip()
            if t:
                lines[line_number] = t
                line_number += 1

    return lines


# * Read Typst (.typ) file as numbered lines
def read_typst(path: Path, preserve_structure: bool = False) -> Lines:
    # ! Lazy import to avoid circular dependency
    from .typst_handler import validate_basic_typst_syntax
    from .typst_patterns import (
        is_preservable_content as typst_is_preservable,
        requires_trailing_blank as typst_requires_trailing_blank,
    )

    try:
        text = Path(path).read_text(encoding="utf-8")
        vlog_file_read(path, len(text))
    except UnicodeDecodeError as e:
        raise TypstError(f"Cannot decode Typst file {path}: {e}")
    except Exception as e:
        raise TypstError(f"Cannot read Typst file {path}: {e}")

    # Validate basic Typst syntax
    if not validate_basic_typst_syntax(text):
        raise TypstError(f"Invalid Typst syntax detected in {path}")

    lines: Lines = {}
    line_number = 1

    if preserve_structure:
        # Structured mode: preserve comments, commands, strategic empty lines
        for raw in text.splitlines():
            t = raw.strip()

            # Preserve important Typst constructs
            if typst_is_preservable(t):
                lines[line_number] = t
                line_number += 1
            # Preserve strategic empty lines after structural commands
            elif raw == "" and line_number > 1:
                prev_line = lines.get(line_number - 1, "")
                if typst_requires_trailing_blank(prev_line):
                    lines[line_number] = ""
                    line_number += 1
    else:
        # Basic mode: strip whitespace, keep only non-empty content
        for raw in text.splitlines():
            t = raw.strip()
            if t:
                lines[line_number] = t
                line_number += 1

    return lines


# * Read resume by file extension (.docx, .tex, or .typ)
def read_resume(path: Path, preserve_structure: bool = False) -> Lines:
    suffix = path.suffix.lower()
    if suffix == ".tex":
        return read_latex(path, preserve_structure=preserve_structure)
    if suffix == ".typ":
        return read_typst(path, preserve_structure=preserve_structure)
    # Use DOCX handling as default
    return read_docx(path)


# * Apply edits to document w/ different preservation modes (in_place: better formatting, rebuild: faster)
def apply_edits_to_docx(
    original_path: Path,
    new_lines: Lines,
    output_path: Path,
    preserve_mode: str = "in_place",
) -> None:
    if preserve_mode == "in_place":
        _apply_edits_in_place(original_path, new_lines, output_path)
    else:
        _apply_edits_rebuild(original_path, new_lines, output_path)


def _apply_edits_in_place(
    original_path: Path, new_lines: Lines, output_path: Path
) -> None:
    # Edit document in-place to preserve formatting & styles
    # Reuse existing reader for lines & paragraph map
    lines, doc, paragraph_map = read_docx_with_formatting(original_path)

    # Compute modifications, additions, deletions
    modifications, additions, deletions = _categorize_edits(lines, new_lines)

    # Apply deletions from end to start
    for line_num in sorted(deletions, reverse=True):
        if line_num in paragraph_map:
            para = paragraph_map[line_num]
            p_element = para._element
            p_element.getparent().remove(p_element)

    # Apply modifications preserving run formatting
    for line_num, new_text in modifications.items():
        if line_num in paragraph_map:
            para = paragraph_map[line_num]
            _set_paragraph_text_preserving_format(para, new_text)

    # Group additions by insertion position
    additions_by_position: Dict[int | None, List[Tuple[int, str]]] = {}
    for insert_after, line_num, text in additions:
        if insert_after not in additions_by_position:
            additions_by_position[insert_after] = []
        additions_by_position[insert_after].append((line_num, text))

    # Sort groups by line number
    for position in additions_by_position:
        additions_by_position[position].sort(key=lambda x: x[0])

    # Insert new paragraph content
    # Sort numeric positions only
    numeric_positions = sorted(
        [p for p in additions_by_position.keys() if p is not None], reverse=True
    )

    # Insert after paragraphs bottom-up
    for insert_after in numeric_positions:
        reference_para = paragraph_map.get(insert_after)
        if not reference_para:
            continue
        for line_num, text in reversed(additions_by_position[insert_after]):
            new_para = _insert_paragraph_after(reference_para, text)
            # Copy reference paragraph style
            if reference_para.style:
                new_para.style = reference_para.style

    # Insert at beginning in reverse order
    if None in additions_by_position:
        for line_num, text in reversed(additions_by_position[None]):
            new_para = doc.add_paragraph(text)
            doc.element.body.insert(0, new_para._element)

    # Persist modified document
    ensure_parent(output_path)
    doc.save(str(output_path))
    vlog_file_write(output_path)


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    # Insert paragraph after given one (python-docx safe)
    new_p = OxmlElement("w:p")
    # Insert after reference paragraph
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _copy_run_formatting(source_run: Run, target_run: Run) -> None:
    # Copy run-level formatting if available
    if not source_run:
        return
    if source_run.font:
        if source_run.font.bold is not None:
            target_run.font.bold = source_run.font.bold
        if source_run.font.italic is not None:
            target_run.font.italic = source_run.font.italic
        if source_run.font.underline is not None:
            target_run.font.underline = source_run.font.underline
        if source_run.font.strike is not None:
            target_run.font.strike = source_run.font.strike
        if source_run.font.size is not None:
            target_run.font.size = source_run.font.size
        if source_run.font.name is not None:
            target_run.font.name = source_run.font.name
        if (
            getattr(source_run.font, "color", None)
            and getattr(source_run.font.color, "rgb", None) is not None
        ):
            target_run.font.color.rgb = source_run.font.color.rgb
        if getattr(source_run.font, "highlight_color", None) is not None:
            target_run.font.highlight_color = source_run.font.highlight_color
    if hasattr(source_run, "style") and source_run.style:
        target_run.style = source_run.style


def _set_paragraph_text_preserving_format(
    target_para: Paragraph, new_text: str, template_para: Paragraph | None = None
) -> None:
    # Set text preserving first-run formatting
    template = template_para if template_para is not None else target_para
    template_run = template.runs[0] if template.runs else None
    target_para.clear()
    new_run = target_para.add_run(new_text)
    if template_run:
        _copy_run_formatting(template_run, new_run)


def _categorize_edits(
    original_lines: Dict[int, str], new_lines: Lines
) -> Tuple[Dict[int, str], List[Tuple[int | None, int, str]], Set[int]]:
    # Categorize edits by type
    modifications: Dict[int, str] = {}
    additions: List[Tuple[int | None, int, str]] = []
    deletions: Set[int] = set()

    original_set = set(original_lines.keys())
    new_set = set(new_lines.keys())

    # Identify modified lines
    for line_num in sorted(new_set):
        if line_num in original_set and new_lines[line_num] != original_lines[line_num]:
            modifications[line_num] = new_lines[line_num]

    # Identify deleted lines
    for line_num in original_set:
        if line_num not in new_set:
            deletions.add(line_num)

    # Identify added lines
    for line_num in sorted(new_set):
        if line_num not in original_set:
            insert_after = None
            for existing_line in sorted(original_set):
                if existing_line < line_num:
                    insert_after = existing_line
                else:
                    break
            additions.append((insert_after, line_num, new_lines[line_num]))

    return modifications, additions, deletions


def _apply_edits_rebuild(
    original_path: Path, new_lines: Lines, output_path: Path
) -> None:
    # Rebuild document from scratch (faster)
    # Load original document
    lines, _, paragraph_map = read_docx_with_formatting(original_path)

    # Sort line numbers for iteration
    sorted_line_nums = sorted(new_lines.keys())
    max_original_line = max(lines.keys()) if lines else 0

    # Track paragraphs & their content
    paragraphs_to_process = []

    for line_num in sorted_line_nums:
        new_text = new_lines[line_num]

        if line_num <= max_original_line and line_num in paragraph_map:
            # Preserve existing paragraph formatting
            para = paragraph_map[line_num]
            paragraphs_to_process.append(("modify", para, new_text))
        else:
            # Create new paragraph
            paragraphs_to_process.append(("new", None, new_text))

    # Create document w/ edits
    new_doc = Document()

    # Note: document styles use defaults

    # Process paragraphs sequentially
    for action, original_para, new_text in paragraphs_to_process:
        if action == "modify" and original_para:
            # Preserve original formatting
            new_para = new_doc.add_paragraph()

            # Copy paragraph formatting
            if original_para.style:
                new_para.style = original_para.style
            if original_para.paragraph_format:
                _copy_paragraph_format(
                    original_para.paragraph_format, new_para.paragraph_format
                )

            # Preserve character formatting via runs
            if original_para.runs:
                _set_paragraph_text_preserving_format(
                    new_para, new_text, template_para=original_para
                )
            else:
                new_para.text = new_text
        else:
            # Create paragraph w/o formatting reference
            new_doc.add_paragraph(new_text)

    # Persist modified document
    ensure_parent(output_path)
    new_doc.save(str(output_path))


def _copy_paragraph_format(source_format: Any, target_format: Any) -> None:
    # Copy paragraph formatting properties
    for prop in dir(source_format):
        if prop.startswith("_"):
            continue
        try:
            value = getattr(source_format, prop)
        except Exception:
            continue
        if callable(value) or value is None:
            continue
        try:
            setattr(target_format, prop, value)
        except Exception:
            # Handle read-only or unsupported properties
            pass


# * Write text to DOCX file (creates plain document)
def write_docx(lines: Lines, output_path: Path) -> None:
    doc = Document()
    for line_num in sorted(lines.keys()):
        doc.add_paragraph(lines[line_num])
    ensure_parent(output_path)
    doc.save(str(output_path))
    vlog_file_write(output_path)


# * Write numbered lines to plain text file (.tex or .txt)
def write_text_lines(lines: Lines, output_path: Path) -> None:
    ordered = "\n".join(f"{text}" for _, text in sorted(lines.items()))
    ensure_parent(output_path)
    Path(output_path).write_text(ordered, encoding="utf-8")
    vlog_file_write(output_path, len(ordered))


# * Read text from file
def read_text(path: Path) -> str:
    text = Path(path).read_text(encoding="utf-8")
    vlog_file_read(path, len(text))
    return text
